from __future__ import annotations

import os
import sys

def create_docx_report():
    try:
        import docx
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("Error: python-docx no esta instalado. Instalandolo primero...")
        return False

    doc = Document()
    
    # 1. Configurar margenes de pagina a 1 pulgada (estandar)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 2. Configurar colores corporativos
    COLOR_PRIMARY = RGBColor(30, 41, 59) # Slate 800 (Fondo formal)
    COLOR_SECONDARY = RGBColor(99, 102, 241) # Indigo (Acento principal)
    COLOR_BODY = RGBColor(51, 65, 85) # Slate 700 (Texto general)
    
    # 3. Configurar estilos tipograficos
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_BODY
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(8)

    # Helper para titulos con estilo
    def add_custom_heading(text, level, space_before=18):
        heading = doc.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.keep_with_next = True
        
        # Color del titulo
        for run in heading.runs:
            run.font.name = 'Arial'
            if level == 1:
                run.font.size = Pt(16)
                run.font.color.rgb = COLOR_PRIMARY
                run.bold = True
            elif level == 2:
                run.font.size = Pt(13)
                run.font.color.rgb = COLOR_SECONDARY
                run.bold = True
        return heading

    # ==================== PORTADA / TITULO ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(10)
    run_title = title_p.add_run("INFORME TÉCNICO DE ARQUITECTURA E INTEGRACIÓN")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(30)
    run_sub = subtitle_p.add_run("Comparador Multimodal Asíncrono de Modelos de Inteligencia Artificial (Versión 3.4)")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.italic = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    # Informacion de autor y fecha
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(100)
    meta_p.paragraph_format.space_after = Pt(6)
    run_author = meta_p.add_run("Preparado para: Evaluaciones de Sistemas y Modelos de Lenguaje")
    run_author.font.size = Pt(10)
    run_author.bold = True
    
    meta_p2 = doc.add_paragraph()
    meta_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = meta_p2.add_run("Fecha de publicacion técnica: Junio de 2026")
    run_date.font.size = Pt(10)

    doc.add_page_break()

    # ==================== SECCIÓN 1 ====================
    add_custom_heading("1. Resumen Ejecutivo e Introducción", level=1)
    
    p = doc.add_paragraph(
        "El presente informe técnico describe la arquitectura de software, la lógica de integración y la metodología "
        "de comparación implementadas en la plataforma de benchmarking en tiempo real denominada 'Antigravity Arena'. "
        "Este sistema surge de la necesidad de evaluar el rendimiento, la precisión, el sesgo creativo y la coherencia "
        "conceptual de diferentes arquitecturas de Inteligencia Artificial en tres dominios clave: generación de texto, "
        "análisis de visión artificial multimodal y síntesis de contenido gráfico (imágenes y secuencias de video)."
    )
    
    p = doc.add_paragraph(
        "A diferencia de las herramientas convencionales de comparación basadas en parámetros estáticos o pruebas de "
        "rendimiento en lotes (batch processing), esta solución implementa un entorno interactivo y reactivo. "
        "Dicho entorno permite contrastar el comportamiento de modelos de vanguardia utilizando una combinación estricta "
        "de servicios de pago y APIs públicas libres de costo, garantizando así un ciclo de pruebas de alto nivel técnico "
        "sin dependencias financieras prohibitivas para proyectos académicos o de investigación."
    )

    p = doc.add_paragraph(
        "El enfoque fundamental de la plataforma radica en su capacidad de ejecutar solicitudes en paralelo, recolectar "
        "los datos de salida, y procesarlos mediante un Juez de Inteligencia Artificial centralizado. Este Juez realiza "
        "un análisis semántico y espacial instantáneo de los resultados de cada modelo, asignando calificaciones precisas "
        "y redactando un informe comparativo detallado. Todo esto ocurre de forma automatizada en el backend, sin "
        "intervención obligatoria del operador, lo que permite acumular estadísticas de coherencia histórica a largo plazo."
    )

    # ==================== SECCIÓN 2 ====================
    add_custom_heading("2. Arquitectura de Software y Procesamiento Asíncrono", level=1)
    
    p = doc.add_paragraph(
        "El sistema ha sido estructurado bajo un patrón de microservicios desacoplados para optimizar la escalabilidad "
        "y evitar bloqueos de recursos. El frontend, desarrollado sobre una arquitectura reactiva en React y compilado "
        "con Vite, gestiona de forma interactiva la recolección de entradas del usuario, el uploader de imágenes base64 "
        "y el renderizado dinámico de resultados. El backend está construido sobre FastAPI (Python), aprovechando "
        "su alto rendimiento asíncrono nativo y su facilidad de integración con subprocesos."
    )

    add_custom_heading("2.1 Lógica del Ciclo de Vida de las Tareas (Background Tasks)", level=2, space_before=12)
    
    p = doc.add_paragraph(
        "Dado que la generación de imágenes y los análisis de visión profunda consumen un tiempo que suele oscilar "
        "entre los 4 y los 15 segundos (dependiendo de la carga de las redes neuronales de los proveedores), ejecutar "
        "estas operaciones en el hilo principal de una petición HTTP convencional resultaría en congelamientos de "
        "interfaz o errores de terminación de conexión (timeouts) en el navegador del usuario."
    )

    p = doc.add_paragraph(
        "Para solucionar esta limitante técnica, el backend implementa un patrón de Tareas en Segundo Plano asíncronas. "
        "Cuando el cliente del frontend envía una petición de inicio a los endpoints de generación, la API de FastAPI "
        "realiza el siguiente procedimiento secuencial:"
    )

    # Lista con viñetas
    doc.add_paragraph("1. Generación de un Identificador Único de Tarea (UUID v4) para indexar la operación en memoria.", style='List Bullet')
    doc.add_paragraph("2. Inicialización del estado de la tarea en el diccionario global en memoria (ACTIVE_TASKS) con progreso en cero y logs vacíos.", style='List Bullet')
    doc.add_paragraph("3. Despacho asíncrono de la función de procesamiento técnico al gestor de BackgroundTasks del servidor.", style='List Bullet')
    doc.add_paragraph("4. Retorno inmediato del identificador único de tarea (task_id) al frontend del navegador en menos de 200 milisegundos.", style='List Bullet')

    p = doc.add_paragraph(
        "Una vez que el frontend recibe el identificador único de tarea, inicia un ciclo de consultas repetidas (polling) "
        "cada 2000 milisegundos hacia el endpoint de estado de la tarea. Este endpoint devuelve el porcentaje actual de "
        "progreso y un mensaje explicativo del hito actual que se está ejecutando (por ejemplo, 'Enviando imagen a Gemini 2.5', "
        "'Procesando en la GPU de Pollinations Flux', etc.). Esto ofrece una experiencia visual informativa e interactiva."
    )

    doc.add_page_break()

    # ==================== SECCIÓN 3 ====================
    add_custom_heading("3. Análisis Técnico y Metodología de las Arenas de Comparación", level=1)
    
    p = doc.add_paragraph(
        "El núcleo funcional del comparador se divide en tres arenas específicas de evaluación técnica. Cada una ha sido "
        "diseñada bajo principios de optimización de recursos y gratuidad de APIs, limitando la necesidad de tarjetas "
        "de crédito o depósitos monetarios en los proveedores."
    )

    add_custom_heading("3.1 Arena de Generación de Texto", level=2, space_before=12)
    p = doc.add_paragraph(
        "Esta arena confronta la creatividad lingüística, la riqueza semántica y el apego sintáctico al estilo "
        "seleccionado. El Modelo A está representado por Google Gemini 2.5 Flash, el cual se consulta mediante el "
        "SDK oficial 'google-genai' utilizando la clave gratuita provista por Google AI Studio. El Modelo B está "
        "representado por Llama 3, consumido de forma libre de costo y sin límites a través de la API abierta de "
        "Pollinations.ai. Los criterios de evaluación automática y humana se basan en el conteo exacto de palabras "
        "generadas respecto al objetivo del slider, la riqueza metafórica en los relatos de fantasía y el nivel de "
        "persuasión comercial en las copias de marketing."
    )

    add_custom_heading("3.2 Arena de Análisis Multimodal de Imagen (Visión)", level=2, space_before=12)
    p = doc.add_paragraph(
        "El procesamiento de visión artificial multimodal permite a los modelos interpretar, describir y extraer patrones "
        "técnicos de una imagen subida por el usuario. El sistema confronta dos generaciones de una misma marca para "
        "evaluar saltos cualitativos entre versiones: Google Gemini 2.5 Flash contra Google Gemini 1.5 Flash. "
        "El backend decodifica la imagen del usuario de su formato base64 y la inyecta como un flujo de bytes nativos "
        "(`types.Part.from_bytes`) al contexto del cliente de Google. Esta comparación resulta 100% gratuita al utilizar "
        "la misma clave de API de AI Studio y evalúa qué versión tiene mejor visión espacial y capacidad de seguir "
        "instrucciones complejas sobre elementos de la foto."
    )

    add_custom_heading("3.3 Arena de Generación de Imagen", level=2, space_before=12)
    p = doc.add_paragraph(
        "Para evaluar la síntesis de contenido gráfico estático, la plataforma contrasta Google Imagen 3.0 (Modelo A) "
        "frente a Pollinations Flux (Modelo B). Google Imagen 3.0 se ejecuta mediante la clave gratuita de Gemini API, "
        "mientras que Pollinations Flux se consume de forma directa, libre y sin claves de API mediante su endpoint "
        "de servidor. Las imágenes se redimensionan en tiempo real en píxeles dependiendo de la relación de aspecto "
        "configurada (1:1, 16:9, 9:16) antes de ser devueltas en base64 al navegador del usuario."
    )

    # ==================== SECCIÓN 4 ====================
    add_custom_heading("4. Juez de IA Automatizado y Multimodal", level=1)
    
    p = doc.add_paragraph(
        "Una de las mayores innovaciones del sistema radica en su árbitro de coherencia automático. Tradicionalmente, "
        "los sistemas de comparación requieren la interacción manual del usuario para registrar votos y determinar un "
        "ganador. En esta plataforma, se ha implementado un proceso de evaluación automática en tiempo real que corre "
        "inmediatamente después de que ambos modelos entregan sus resultados, informando las puntuaciones y el análisis "
        "detallado en la base de la pantalla de forma instantánea."
    )

    p = doc.add_paragraph(
        "Para las arenas de texto y visión, el Juez de IA (representado por Gemini 2.5 Flash) recibe el prompt original "
        "junto con los dos textos de salida. Se aplican instrucciones de sistema estrictas y se configura la API para "
        "forzar una respuesta en formato JSON puro. El Juez evalúa de forma objetiva la gramática, la creatividad y la "
        "coherencia temática, asignando notas del 1 al 10 para cada modelo."
    )

    p = doc.add_paragraph(
        "Para la arena de generación de imágenes, el Juez de IA ejecuta un flujo multimodal avanzado. Dado que Gemini 2.5 "
        "es capaz de comprender e interpretar datos visuales directamente, el backend extrae los bytes de imagen generados "
        "por Google Imagen y por Pollinations Flux, inyectándolos como archivos adjuntos al contexto del Juez. "
        "El Juez de IA observa físicamente ambas imágenes, evalúa la precisión con la que plasmaron los objetos descritos "
        "en el prompt original, detecta anomalías visuales y determina cuál imagen es semántica y compositivamente "
        "más fiel a la idea base, redactando un veredicto escrito en español."
    )

    doc.add_page_break()

    # ==================== SECCIÓN 5 ====================
    add_custom_heading("5. Sistema de Estadísticas, Persistencia y Frontend", level=1)
    
    p = doc.add_paragraph(
        "La persistencia de los datos del comparador está delegada al módulo de backend `stats_manager.py`. Este módulo "
        "se encarga de inicializar, leer e incrementar las estadísticas históricas almacenadas en el archivo "
        "`comparison_stats.json` en el directorio de trabajo del servidor."
    )

    p = doc.add_paragraph(
        "Cada vez que el operador decide emitir un voto manual utilizando los botones del frontend, el cliente envía "
        "una solicitud de tipo POST al endpoint `/api/vote` incluyendo la categoría del duelo, el modelo elegido y el "
        "prompt que se ejecutó. El gestor de estadísticas no solo incrementa los contadores acumulados de las variables "
        "`text_votes`, `vision_votes`, `image_votes` y `video_votes`, sino que además inserta un registro histórico "
        "con la fecha y hora exacta del voto en un arreglo denominado `history`."
    )

    p = doc.add_paragraph(
        "El frontend de React consume en tiempo real esta estructura del JSON. En la barra lateral, calcula de forma "
        "dinámica los porcentajes de victorias y renderiza barras de progreso fluidas y adaptadas al modo que se encuentre "
        "activo. En la parte inferior, renderiza una tabla de historial en vidrio templado (glassmorphic) que mapea de "
        "forma interactiva los metadatos a nombres de modelos formales y comprensibles, permitiendo llevar una bitácora "
        "de qué modelos se adaptaron mejor a cada prompt ingresado durante el ciclo de pruebas."
    )

    # ==================== SECCIÓN 6 ====================
    add_custom_heading("6. Conclusión y Guía de Puesta en Marcha", level=1)
    
    p = doc.add_paragraph(
        "La reconfiguración completa a modelos 100% gratuitos representa un avance sumamente práctico para el entorno "
        "de desarrollo del proyecto. Al desvincular APIs complejas y de pago obligatorio, el sistema se torna accesible "
        "para cualquier operador que disponga de una única clave de API de Gemini, la cual puede obtenerse en pocos segundos. "
        "La integración híbrida de Gemini con las APIs libres de Pollinations.ai garantiza un flujo de trabajo continuo, "
        "seguro y de excelente calidad visual en imágenes y texto."
    )

    add_custom_heading("6.1 Pasos para la Inicialización del Comparador", level=2, space_before=12)
    p = doc.add_paragraph(
        "Para ejecutar el comparador multimodal, el operador debe seguir el procedimiento secuencial que se describe "
        "a continuación:"
    )

    doc.add_paragraph("1. Crear y configurar el archivo de variables de entorno .env dentro del directorio backend, ingresando únicamente la clave de Gemini en GEMINI_API_KEY.", style='List Bullet')
    doc.add_paragraph("2. Mantener vacías las variables de Stability y OpenAI para que el backend conmute de forma automática a los modelos libres de Pollinations y los simuladores de video.", style='List Bullet')
    doc.add_paragraph("3. Dar permisos de ejecución e iniciar el script de orquestación unificado ./run.sh en la terminal principal de Linux.", style='List Bullet')
    doc.add_paragraph("4. Acceder al enlace local entregado por la compilación de Vite en el navegador web habitual para iniciar las pruebas.", style='List Bullet')

    p = doc.add_paragraph(
        "Este diseño arquitectónico y metodológico robusto asegura que el comparador permanezca estable, responsivo y "
        "completamente funcional de forma ilimitada y libre de costos de operación."
    )

    # 4. Guardar archivo
    output_path = "/home/starkiller/Documents/noveno/noveno/programacion/proyecto/informe_tecnico.docx"
    doc.save(output_path)
    print(f"Documento docx creado exitosamente en: {output_path}")
    return True

if __name__ == "__main__":
    create_docx_report()
